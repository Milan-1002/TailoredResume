from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument

from models import ResumeBullet, TailoredResume

OUTPUT_DIR = Path(__file__).parent / "output"
DEFAULT_OUTPUT_PATH = OUTPUT_DIR / "tailored_resume.docx"


@dataclass
class DocxExportResult:
    path: Path
    edited_count: int
    skipped_count: int


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Keeps the paragraph's existing formatting by reusing its first run
    (font, size, bold, etc.) and just swapping the text — this is what
    preserves the original document's look instead of rebuilding it."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def export_tailored_docx(
    tailored: TailoredResume,
    all_bullets: list[ResumeBullet],
    original_docx_path: Path,
    output_path: Path = DEFAULT_OUTPUT_PATH,
) -> DocxExportResult:
    bullets_by_id = {b.id: b for b in all_bullets}
    doc = DocxDocument(str(original_docx_path))
    paragraphs = doc.paragraphs

    edited_count = 0
    skipped_count = 0

    for rb in tailored.bullets:
        source_bullet = bullets_by_id.get(rb.bullet_id)
        if source_bullet is None or source_bullet.source_paragraph_index is None:
            skipped_count += 1
            continue
        if source_bullet.source_paragraph_index >= len(paragraphs):
            skipped_count += 1
            continue
        if rb.text.strip() == source_bullet.text.strip():
            continue

        _replace_paragraph_text(paragraphs[source_bullet.source_paragraph_index], rb.text)
        edited_count += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return DocxExportResult(path=output_path, edited_count=edited_count, skipped_count=skipped_count)
