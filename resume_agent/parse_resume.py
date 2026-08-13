import difflib
import re
import sys
import uuid
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from pydantic import BaseModel

from db import save_bullets
from embeddings import embed_texts
from llm import call_structured, run_concurrent
from models import ResumeBullet
from prompts import RESUME_EXTRACTION_SYSTEM_PROMPT, RESUME_SEGMENTATION_SYSTEM_PROMPT

FUZZY_MATCH_THRESHOLD = 0.85


class ResumeSection(BaseModel):
    employer: str
    role: str
    start_date: str
    end_date: str | None
    raw_bullets_text: str


class ResumeSections(BaseModel):
    sections: list[ResumeSection]


class ExtractedBullet(BaseModel):
    text: str
    skills: list[str]
    seniority_signal: str | None
    quantified_outcome: str | None


class ExtractedBullets(BaseModel):
    bullets: list[ExtractedBullet]


def read_resume_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if suffix == ".pdf":
        chunks = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    chunks.append(text)
        return "\n".join(chunks)
    return path.read_text()


def segment_resume(text: str) -> list[ResumeSection]:
    result = call_structured(
        system=RESUME_SEGMENTATION_SYSTEM_PROMPT,
        user=f"Resume text:\n\n{text}",
        schema=ResumeSections,
        tool_name="record_resume_sections",
    )
    return result.sections


def extract_bullets_for_section(section: ResumeSection) -> list[ResumeBullet]:
    user = (
        f"Employer: {section.employer}\n"
        f"Role: {section.role}\n"
        f"Dates: {section.start_date} - {section.end_date or 'present'}\n\n"
        f"Raw bullet text:\n{section.raw_bullets_text}"
    )
    result = call_structured(
        system=RESUME_EXTRACTION_SYSTEM_PROMPT,
        user=user,
        schema=ExtractedBullets,
        tool_name="record_bullets",
    )
    return [
        ResumeBullet(
            id=str(uuid.uuid4()),
            employer=section.employer,
            role=section.role,
            start_date=section.start_date,
            end_date=section.end_date,
            text=b.text,
            skills=b.skills,
            seniority_signal=b.seniority_signal,
            quantified_outcome=b.quantified_outcome,
        )
        for b in result.bullets
    ]


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().lower()


def match_bullets_to_paragraphs(bullets: list[ResumeBullet], docx_path: Path) -> None:
    """Finds which original .docx paragraph each bullet came from, so
    docx_export.py can later edit that exact paragraph in place instead of
    rebuilding the document from scratch. Sets source_paragraph_index in place;
    leaves it None where no confident match is found (that bullet simply won't
    be edited during export)."""
    doc = DocxDocument(str(docx_path))
    paragraphs = [(i, _normalize(p.text)) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    used_indices: set[int] = set()

    for bullet in bullets:
        normalized_bullet = _normalize(bullet.text)

        exact = next(
            (i for i, ptext in paragraphs if i not in used_indices and ptext == normalized_bullet),
            None,
        )
        if exact is not None:
            bullet.source_paragraph_index = exact
            used_indices.add(exact)
            continue

        best_index, best_ratio = None, 0.0
        for i, ptext in paragraphs:
            if i in used_indices:
                continue
            ratio = difflib.SequenceMatcher(None, ptext, normalized_bullet).ratio()
            if ratio > best_ratio:
                best_index, best_ratio = i, ratio

        if best_index is not None and best_ratio >= FUZZY_MATCH_THRESHOLD:
            bullet.source_paragraph_index = best_index
            used_indices.add(best_index)


def parse_resume(path: Path) -> list[ResumeBullet]:
    text = read_resume_text(path)
    sections = segment_resume(text)

    # Each section (job) is extracted independently, so run these concurrently
    # instead of one-at-a-time — cuts upload time roughly proportional to the
    # number of jobs on the resume.
    bullets: list[ResumeBullet] = []
    for section_bullets in run_concurrent(extract_bullets_for_section, sections):
        bullets.extend(section_bullets)

    if path.suffix.lower() == ".docx":
        match_bullets_to_paragraphs(bullets, path)

    if bullets:
        vectors = embed_texts([b.text for b in bullets])
        for bullet, vector in zip(bullets, vectors):
            bullet.embedding = vector

    return bullets


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python parse_resume.py path/to/master_resume.docx", file=sys.stderr)
        sys.exit(1)

    resume_path = Path(sys.argv[1])
    parsed_bullets = parse_resume(resume_path)
    save_bullets(parsed_bullets)

    print(f"Parsed {len(parsed_bullets)} bullets from {resume_path}")
    for b in parsed_bullets:
        print(f"- [{b.employer} / {b.role}] {b.text[:90]}")
